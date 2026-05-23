from docx import Document
from docx.oxml import OxmlElement

src = r"C:\Users\zhu20\Desktop\毕业论文\毕设论文_在线电影购票系统_终稿完整版.docx"
out = r"C:\Users\zhu20\Desktop\毕业论文\毕设论文_在线电影购票系统_终稿完整版_v2_全替换表设计.docx"

doc = Document(src)

# 1) 全文术语替换（仅做较稳妥替换）
replacements = {
    '酒店预订': '电影购票',
    '酒店预约': '电影购票',
    '酒店订房': '电影购票',
    '酒店管理': '影院管理',
    '酒店信息': '影院信息',
    '酒店': '影院',
    '客房': '影厅',
    '房型': '影厅类型',
    '入住': '观影',
    '退房': '退票',
    '订房': '购票',
    '预订': '购票',
}

replace_hits = 0
for p in doc.paragraphs:
    t = p.text
    if not t:
        continue
    nt = t
    for k, v in replacements.items():
        if k in nt:
            nt = nt.replace(k, v)
    if nt != t:
        replace_hits += 1
        p.text = nt

# 2) 定位“4.3 数据库…”到“4.4 ...”并重写为电影购票表设计
start = -1
end = -1
for i, p in enumerate(doc.paragraphs):
    t = (p.text or '').strip()
    if start == -1 and (t.startswith('4.3') and ('数据库' in t or '数据表' in t)):
        start = i
    elif start != -1 and t.startswith('4.4'):
        end = i
        break

if start == -1:
    # 兜底：找第4章位置后插入
    for i,p in enumerate(doc.paragraphs):
        if (p.text or '').strip().startswith('第4章'):
            start = i + 6
            end = start + 1
            break

if end == -1:
    end = min(len(doc.paragraphs)-1, start + 8)

block = [
"4.3 数据库表设计（电影购票系统）",
"本系统采用 MySQL 存储核心业务数据，围绕“电影—影院—场次—座位—订单”主链路进行建模，同时补充用户互动与后台管理数据表。",
"4.3.1 users（用户表）",
"字段设计：_id（主键，用户唯一标识）、_openid（微信标识）、phone（手机号）、nickName（昵称）、avatarUrl（头像）、isAdmin（管理员标记）、createTime、updateTime。",
"约束说明：_id 为主键；phone 建议建立普通索引；isAdmin 默认 0，用于后台权限判断。",
"4.3.2 movies（电影表）",
"字段设计：_id（主键，tmdb_xxx）、title（片名）、poster（海报地址）、rating（评分）、genre（类型）、duration（片长）、director（导演）、actors（主演）、description（简介）、releaseDate（上映日期）、price（基础票价，分）、status（showing/coming/off）、hot（热度）、createTime、updateTime。",
"约束说明：_id 主键；status 建议索引；hot 用于排序展示。",
"4.3.3 cinemas（影院表）",
"字段设计：_id（主键）、name（影院名）、address（地址）、phone（电话）、latitude、longitude、city（城市）、district（区县）、minPrice（最低价，分）、tags（标签）、facilities（设施）、distance（距离缓存）、createTime、updateTime。",
"约束说明：_id 主键；city 建议索引，支持按城市筛选。",
"4.3.4 schedules（场次表）",
"字段设计：_id（主键，movieId_cinemaId_date_index）、movieId、cinemaId、hallName、hallType、date、startTime、endTime、price（分）、totalSeats、availableSeats、status、createTime、updateTime。",
"约束说明：_id 主键；cinemaId+date 建议联合索引；movieId 建议索引。",
"4.3.5 seats（座位表）",
"字段设计：_id（主键，scheduleId_row_col）、scheduleId、rowNum、colNum、status（available/locked/sold）、orderId（占用该座位的订单）、createTime、updateTime。",
"约束说明：scheduleId 建议索引；scheduleId+rowNum+colNum 建议唯一约束，防止重复座位。",
"4.3.6 orders（订单表）",
"字段设计：_id（主键）、orderNo（订单号）、_openid（用户标识）、scheduleId、movieId、cinemaId、movieTitle、moviePoster、cinemaName、hallName、date、startTime、seatsJson（座位集合 JSON）、seatCount、totalPrice（分）、couponId、status（pending/paid/cancelled/refunded）、purchaseTime、payTime、createTime、updateTime。",
"约束说明：_id 主键；_openid+status 建议索引；orderNo 建议唯一索引。",
"4.3.7 collections（收藏表）",
"字段设计：_id（主键）、_openid、movieId、title、poster、createTime。",
"约束说明：_openid+movieId 唯一约束，防止重复收藏。",
"4.3.8 movie_comments（评论表）",
"字段设计：_id（主键）、movieId、_openid、nickName、avatarUrl、rating、content、likes、createTime。",
"约束说明：movieId 建议索引，支持评论列表分页查询。",
"4.3.9 表关系与一致性说明",
"关系链路：movies 1..n schedules；cinemas 1..n schedules；schedules 1..n seats；orders 与 schedules/seats 通过 scheduleId/orderId 关联；collections、movie_comments 通过 movieId 与 movies 关联。",
"一致性策略：下单时在事务中锁座，支付后改 sold，取消/退款回滚 available，确保订单与座位状态一致。",
"4.3.10 索引与性能建议",
"建议建立索引：movies(status, hot)、cinemas(city)、schedules(cinemaId, date)、schedules(movieId)、orders(_openid, status)、seats(scheduleId)。在千级电影与多影院场景下可显著提升查询效率。",
]

# 清空旧段落区间
for i in range(start, end):
    doc.paragraphs[i].text = ''

# 在 start 位置写入第一行
if start < len(doc.paragraphs):
    doc.paragraphs[start].text = block[0]
    anchor = doc.paragraphs[start]
else:
    anchor = doc.add_paragraph(block[0])

# 在 anchor 后插入段落
for line in block[1:]:
    new_p = OxmlElement('w:p')
    anchor._p.addnext(new_p)
    para = anchor._parent.add_paragraph('')
    # 将新段落移动到 new_p 位置
    new_p.getparent().replace(new_p, para._p)
    para.text = line
    anchor = para

# 3) 再次扫描是否残留酒店词（统计）
left_keywords = ['酒店', '客房', '入住', '订房', '预订']
left_hits = {k:0 for k in left_keywords}
for p in doc.paragraphs:
    t = p.text or ''
    for k in left_keywords:
        if k in t:
            left_hits[k] += 1

doc.save(out)
print(out)
print('replace_hits=', replace_hits)
print('left_hits=', left_hits)
print('section_start=', start, 'section_end=', end)
