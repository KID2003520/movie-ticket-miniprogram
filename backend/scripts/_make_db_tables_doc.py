from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

src = r"C:\Users\zhu20\Desktop\毕业论文\毕设论文_在线电影购票系统_终稿完整版_v2_全替换表设计.docx"
out = r"C:\Users\zhu20\Desktop\毕业论文\毕设论文_在线电影购票系统_终稿完整版_v3_数据库三线表.docx"

doc = Document(src)

# 找到 4.3 起止
start = end = -1
for i,p in enumerate(doc.paragraphs):
    t=(p.text or '').strip()
    if start==-1 and t.startswith('4.3') and ('数据库' in t or '数据表' in t):
        start=i
    elif start!=-1 and t.startswith('4.4'):
        end=i
        break

if start==-1:
    raise SystemExit('未找到 4.3 章节')
if end==-1:
    end=len(doc.paragraphs)

# 清空 4.3~4.4 之间旧文本
for i in range(start, end):
    doc.paragraphs[i].text = ''

doc.paragraphs[start].text = '4.3 数据库表设计（标准表格版）'
anchor = doc.paragraphs[start]


def add_para_after(anchor_para, text):
    new_p = OxmlElement('w:p')
    anchor_para._p.addnext(new_p)
    para = anchor_para._parent.add_paragraph('')
    new_p.getparent().replace(new_p, para._p)
    para.text = text
    return para

def add_table_after(anchor_para, rows_data):
    tbl = doc.add_table(rows=1, cols=4)
    hdr = tbl.rows[0].cells
    hdr[0].text = '字段名'
    hdr[1].text = '类型'
    hdr[2].text = '约束/索引'
    hdr[3].text = '说明'
    for r in rows_data:
        cells = tbl.add_row().cells
        for i,v in enumerate(r):
            cells[i].text = v

    # 简单三线表样式：使用 Table Grid 并去掉中间竖线很复杂，这里采用学校普遍可接受的网格表，便于后续手工一键改样式
    tbl.style = 'Table Grid'

    # 把表移动到锚点后
    new_tbl = OxmlElement('w:tbl')
    anchor_para._p.addnext(new_tbl)
    new_tbl.getparent().replace(new_tbl, tbl._tbl)

    # 返回表后新段落锚点
    p = add_para_after(anchor_para, '')
    return p

# 总述
anchor = add_para_after(anchor, '本系统核心数据表包括 users、movies、cinemas、schedules、seats、orders、collections、movie_comments。下表给出各表关键字段设计。')

schemas = [
    ('表4-1 users（用户表）', [
        ('_id', 'VARCHAR(64)', 'PK', '用户主键ID'),
        ('_openid', 'VARCHAR(64)', 'INDEX', '微信标识'),
        ('phone', 'VARCHAR(20)', 'INDEX', '手机号'),
        ('nickName', 'VARCHAR(64)', '', '昵称'),
        ('avatarUrl', 'VARCHAR(512)', '', '头像地址'),
        ('isAdmin', 'TINYINT', 'DEFAULT 0', '管理员标识'),
        ('createTime', 'DATETIME', '', '创建时间'),
        ('updateTime', 'DATETIME', '', '更新时间'),
    ]),
    ('表4-2 movies（电影表）', [
        ('_id', 'VARCHAR(64)', 'PK', '电影主键（tmdb_xxx）'),
        ('title', 'VARCHAR(255)', 'INDEX', '电影标题'),
        ('poster', 'VARCHAR(512)', '', '海报URL'),
        ('rating', 'DECIMAL(3,1)', '', '评分'),
        ('genre', 'VARCHAR(128)', '', '类型'),
        ('duration', 'INT', '', '片长（分钟）'),
        ('director', 'VARCHAR(255)', '', '导演'),
        ('actors', 'TEXT', '', '主演'),
        ('description', 'TEXT', '', '剧情简介'),
        ('releaseDate', 'DATE', '', '上映日期'),
        ('price', 'INT', '', '基础票价（分）'),
        ('status', 'VARCHAR(16)', 'INDEX', 'showing/coming/off'),
        ('hot', 'INT', 'INDEX', '热度值'),
    ]),
    ('表4-3 cinemas（影院表）', [
        ('_id', 'VARCHAR(64)', 'PK', '影院主键ID'),
        ('name', 'VARCHAR(255)', 'INDEX', '影院名称'),
        ('address', 'VARCHAR(255)', '', '影院地址'),
        ('phone', 'VARCHAR(32)', '', '联系电话'),
        ('latitude', 'DECIMAL(10,6)', '', '纬度'),
        ('longitude', 'DECIMAL(10,6)', '', '经度'),
        ('city', 'VARCHAR(64)', 'INDEX', '城市'),
        ('minPrice', 'INT', '', '最低票价（分）'),
        ('tags', 'VARCHAR(255)', '', '标签'),
    ]),
    ('表4-4 schedules（场次表）', [
        ('_id', 'VARCHAR(128)', 'PK', '场次主键ID'),
        ('movieId', 'VARCHAR(64)', 'INDEX', '关联电影ID'),
        ('cinemaId', 'VARCHAR(64)', 'INDEX', '关联影院ID'),
        ('hallName', 'VARCHAR(64)', '', '影厅名称'),
        ('hallType', 'VARCHAR(64)', '', '影厅类型'),
        ('date', 'DATE', 'INDEX', '放映日期'),
        ('startTime', 'VARCHAR(16)', '', '开始时间'),
        ('endTime', 'VARCHAR(16)', '', '结束时间'),
        ('price', 'INT', '', '场次票价（分）'),
        ('totalSeats', 'INT', '', '总座位数'),
        ('availableSeats', 'INT', '', '可售座位数'),
        ('status', 'VARCHAR(16)', '', '场次状态'),
    ]),
    ('表4-5 seats（座位表）', [
        ('_id', 'VARCHAR(128)', 'PK', '座位记录ID'),
        ('scheduleId', 'VARCHAR(128)', 'INDEX', '关联场次ID'),
        ('rowNum', 'INT', '', '行号'),
        ('colNum', 'INT', '', '列号'),
        ('status', 'VARCHAR(16)', 'INDEX', 'available/locked/sold'),
        ('orderId', 'VARCHAR(64)', 'INDEX', '关联订单ID'),
    ]),
    ('表4-6 orders（订单表）', [
        ('_id', 'VARCHAR(64)', 'PK', '订单主键ID'),
        ('orderNo', 'VARCHAR(64)', 'UNIQUE', '订单编号'),
        ('_openid', 'VARCHAR(64)', 'INDEX', '用户标识'),
        ('scheduleId', 'VARCHAR(128)', 'INDEX', '场次ID'),
        ('movieId', 'VARCHAR(64)', 'INDEX', '电影ID'),
        ('cinemaId', 'VARCHAR(64)', 'INDEX', '影院ID'),
        ('seatsJson', 'TEXT', '', '座位JSON'),
        ('seatCount', 'INT', '', '座位数量'),
        ('totalPrice', 'INT', '', '订单金额（分）'),
        ('status', 'VARCHAR(16)', 'INDEX', 'pending/paid/cancelled/refunded'),
        ('purchaseTime', 'DATETIME', '', '下单时间'),
        ('payTime', 'DATETIME', '', '支付时间'),
    ]),
    ('表4-7 collections（收藏表）', [
        ('_id', 'VARCHAR(64)', 'PK', '收藏主键ID'),
        ('_openid', 'VARCHAR(64)', 'INDEX', '用户标识'),
        ('movieId', 'VARCHAR(64)', 'INDEX', '电影ID'),
        ('title', 'VARCHAR(255)', '', '电影标题快照'),
        ('poster', 'VARCHAR(512)', '', '海报快照'),
        ('createTime', 'DATETIME', '', '收藏时间'),
    ]),
    ('表4-8 movie_comments（评论表）', [
        ('_id', 'VARCHAR(64)', 'PK', '评论主键ID'),
        ('movieId', 'VARCHAR(64)', 'INDEX', '电影ID'),
        ('_openid', 'VARCHAR(64)', 'INDEX', '用户标识'),
        ('nickName', 'VARCHAR(64)', '', '昵称'),
        ('avatarUrl', 'VARCHAR(512)', '', '头像'),
        ('rating', 'INT', '', '评分'),
        ('content', 'TEXT', '', '评论内容'),
        ('likes', 'INT', 'DEFAULT 0', '点赞数'),
        ('createTime', 'DATETIME', '', '创建时间'),
    ]),
]

for title, rows in schemas:
    anchor = add_para_after(anchor, title)
    anchor = add_table_after(anchor, rows)

anchor = add_para_after(anchor, '4.3.11 表关系说明：movies 与 schedules 为一对多；cinemas 与 schedules 为一对多；schedules 与 seats 为一对多；orders 通过 scheduleId/orderId 与场次、座位关联；collections、movie_comments 通过 movieId 与电影关联。')
anchor = add_para_after(anchor, '4.3.12 一致性说明：下单时先锁座再落单，支付后置 sold，取消/退款回滚 available，确保订单与座位状态一致。')

# 删除 start~end 间空段落太多不强求，保留即可

doc.save(out)
print(out)
