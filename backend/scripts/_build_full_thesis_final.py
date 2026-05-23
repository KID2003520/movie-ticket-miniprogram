from docx import Document

src = r"C:\Users\zhu20\Desktop\毕业论文\毕设论文参考模板2026.docx"
out = r"C:\Users\zhu20\Desktop\毕业论文\毕设论文_在线电影购票系统_终稿完整版.docx"

doc = Document(src)

# 封面与摘要关键位置覆盖（沿用模板版式）
updates = {
    12: "题目：基于微信小程序的在线电影购票系统设计与实现",
    29: "毕业设计题目：基于微信小程序的在线电影购票平台",
    33: "本课题围绕“基于微信小程序的在线电影购票系统”开展设计与实现，构建一套覆盖用户端、管理端与数据运维端的完整票务业务系统。系统实现电影浏览、影院与场次查询、在线选座、订单创建/取消/支付/退款、收藏评论、管理端统计与运维工具等功能。",
    65: "摘  要",
    66: "随着移动互联网和在线支付的发展，电影票务业务逐步由线下窗口迁移至线上平台。传统购票方式在信息透明度、购票效率和服务体验上存在明显不足。本文以“在线电影购票微信小程序”为研究对象，完成了系统需求分析、系统架构设计、数据库设计、功能实现与测试验证。系统前端采用微信小程序原生框架，后端采用 Node.js 与 Express 构建 REST 服务，数据层采用 MySQL 持久化存储，并支持 TMDB 真实电影数据批量导入。系统核心功能包括电影检索、影院筛选、场次查询、可视化选座、订单生命周期管理、评论收藏、管理员后台管理等。针对并发购票场景，系统采用事务化锁座机制，保证座位状态与订单状态一致，避免重复售卖。测试结果表明，系统在主要业务链路上运行稳定，能够满足毕业设计场景下的功能需求与演示需求。",
    67: "关键词：在线电影购票；微信小程序；Node.js；Express；MySQL；选座系统",
    69: "ABSTRACT",
    70: "This thesis presents the design and implementation of an online movie ticketing system based on WeChat Mini Program. The system addresses common issues in traditional ticket purchasing, such as low efficiency and poor transparency. The frontend is implemented with native mini-program technologies, while the backend is built with Node.js and Express, and MySQL is used for persistent storage. Core functions include movie browsing, cinema and schedule query, visual seat selection, order lifecycle management, collection and comments, and admin operations. A transaction-based seat-locking strategy is adopted to ensure consistency between seat status and order status in concurrent scenarios. The system also supports TMDB batch import for real movie data. Experimental results show that the system is stable, practical, and extensible for graduation project applications.",
    71: "Key words: online movie ticketing; WeChat Mini Program; Node.js; Express; MySQL; seat selection"
}
for i, t in updates.items():
    if i < len(doc.paragraphs):
        doc.paragraphs[i].text = t

# 正文全量重写（从目录后正文开始）
start_idx = 74
content = [
"第1章 绪论",
"1.1 研究背景",
"近年来，电影消费场景持续线上化，用户对“随时可查、随时可购、可视化选座”的体验要求不断提高。传统线下购票方式存在排队时间长、场次更新不及时、座位不可视等问题，难以满足当前用户需求。微信小程序凭借无需安装、入口便捷、与支付生态天然衔接等优势，成为票务系统轻量化落地的重要技术载体。",
"1.2 研究意义",
"本课题从工程实践角度构建完整在线电影购票系统，具有以下意义：第一，验证微信小程序在票务业务中的可行性；第二，形成可复用的前后端分离实现方案；第三，为后续接入真实支付、推荐算法与运营体系提供可扩展基础。",
"1.3 国内外研究现状",
"国内主流票务平台已形成较成熟的线上购票能力，但高校项目与中小系统普遍存在数据来源不稳定、并发一致性不足、运维脚本缺失等问题。国外在线电影服务在推荐和智能运营方面发展较快，但其实现方案难以直接迁移到微信生态。基于此，本文聚焦“可实现、可演示、可维护”的工程目标。",
"1.4 研究内容与论文结构",
"本文围绕需求分析、系统设计、系统实现、系统测试四条主线展开。全文结构如下：第1章为绪论；第2章介绍关键技术；第3章给出需求分析；第4章进行系统设计；第5章介绍系统实现；第6章给出测试分析；最后给出结论与展望。",

"第2章 关键技术与开发环境",
"2.1 微信小程序原生开发技术",
"系统前端采用 WXML + WXSS + JavaScript。WXML 负责页面结构，WXSS 负责样式，JavaScript 负责业务逻辑与数据交互。利用小程序生命周期与本地缓存机制，实现页面状态管理与用户体验优化。",
"2.2 Node.js 与 Express",
"后端采用 Express 构建 REST API 服务。通过模块化路由设计将电影、影院、场次、订单、评论收藏、管理等功能拆分为独立服务。Express 中间件用于统一处理跨域、JSON 解析、错误返回与日志输出。",
"2.3 MySQL 数据库",
"系统采用 MySQL 存储核心业务数据，包含 movies、cinemas、schedules、seats、orders、users、collections、movie_comments 等表。通过主键、外键逻辑与索引设计实现业务关联与查询性能保障。",
"2.4 TMDB 数据接入技术",
"系统通过 TMDB API 获取真实电影数据，支持多榜单导入与 discover 批量导入。为适配网络环境，支持 HTTPS_PROXY 代理配置。导入后通过状态重算脚本自动更新影片 showing/coming 状态。",
"2.5 开发与运行环境",
"开发环境为 Windows + 微信开发者工具 + Node.js + MySQL。后端默认监听 3000 端口，前端通过配置项连接后端服务地址。",

"第3章 系统需求分析",
"3.1 需求分析目标",
"需求分析目标是明确系统角色、业务边界、功能清单和质量要求，为后续设计与实现提供完整依据。",
"3.2 角色分析",
"系统包含三类角色：普通用户、管理员、运维人员。普通用户完成购票业务；管理员负责内容与运营管理；运维人员负责数据导入、排期重建和服务维护。",
"3.3 用户端功能需求",
"3.3.1 登录与身份需求：支持登录态保持、用户信息读取与登录拦截。",
"3.3.2 电影浏览需求：支持热映/待映分类、关键词搜索、详情查看。",
"3.3.3 影院场次需求：支持城市筛选、日期切换、电影维度过滤。",
"3.3.4 选座下单需求：支持可视化座位图、价格校验、订单创建。",
"3.3.5 订单管理需求：支持待支付、已支付、已取消、已退款状态流转。",
"3.3.6 互动需求：支持收藏、取消收藏、评论发布与评论展示。",
"3.4 管理端需求",
"管理员可执行电影上下架、电影删除、用户查看、统计查看、排期重建、导入任务触发等操作。",
"3.5 数据需求",
"关键实体包括用户、电影、影院、场次、座位、订单、评论、收藏。movieId、cinemaId、scheduleId、orderId 为核心关联键。",
"3.6 业务流程需求",
"购票主流程：电影详情→影院选择→场次选择→选座→创建订单→支付成功。",
"退款流程：支付订单发起退款→订单状态更新→座位状态回滚。",
"异常流程：座位冲突、参数错误、网络失败时应保持数据一致并返回明确错误。",
"3.7 非功能需求",
"性能：核心接口响应可接受；可靠性：订单与座位状态一致；安全性：接口权限校验与参数校验；可维护性：模块分层与脚本化运维。",
"3.8 本章小结",
"本章明确了系统在功能与质量层面的建设目标，为第4章系统设计奠定基础。",

"第4章 系统设计",
"4.1 总体架构设计",
"系统采用前后端分离架构：小程序前端负责展示与交互，Express 后端负责业务处理，MySQL 负责数据持久化。通过统一 API 协议减少耦合，提高维护效率。",
"4.2 模块划分设计",
"前端模块：首页、电影页、电影详情页、影院页、影院详情页、选座页、订单页、订单详情页、用户中心、管理页。",
"后端模块：电影模块、影院模块、场次模块、订单模块、评论收藏模块、管理模块、位置服务模块、TMDB 导入模块。",
"4.3 数据库结构设计",
"movies：影片基础信息与展示字段；",
"cinemas：影院名称、地址、经纬度与标签；",
"schedules：场次时间、影厅、价格与可售座位；",
"seats：按场次维度维护座位状态；",
"orders：订单主表及状态、金额、座位信息；",
"users：用户基础资料与权限字段；",
"collections/movie_comments：互动数据。",
"4.4 关键机制设计",
"4.4.1 锁座机制：创建订单前将座位从 available 更新为 locked，失败则回滚。",
"4.4.2 订单状态机：pending→paid/cancelled/refunded，保证状态转移合法。",
"4.4.3 数据导入机制：批量导入电影后重算状态并支持排期重建。",
"4.5 接口设计原则",
"接口统一返回 code/message/data 结构；关键接口进行参数校验与身份校验；错误信息明确可追踪。",
"4.6 本章小结",
"本章完成了系统架构、模块、数据库与核心机制设计，为系统编码实现提供了结构化指导。",

"第5章 系统实现",
"5.1 前端实现",
"5.1.1 首页与电影列表：支持热映/待映切换与分页展示。",
"5.1.2 电影详情：展示影片完整信息，支持收藏、评论和购票入口。",
"5.1.3 影院与场次：支持城市筛选、定位排序、按日期切换场次。",
"5.1.4 选座页面：展示座位矩阵并校验可选状态。",
"5.1.5 订单页面：展示订单状态并提供支付、取消、删除等操作。",
"5.2 后端实现",
"5.2.1 电影接口：电影列表、详情、搜索、TMDB 补全。",
"5.2.2 影院场次接口：影院列表、影院详情、场次查询与按需补排。",
"5.2.3 订单接口：创建订单、取消订单、模拟支付、退款、删除。",
"5.2.4 管理接口：用户列表、仪表盘、电影状态更新、排期重建。",
"5.3 脚本与运维实现",
"实现了 importAllMoviesTmdb、importDiscoverBulk、repairMoviesAndSchedules 等脚本，支持批量导入真实数据与快速重建演示环境。",
"5.4 关键实现说明",
"在订单创建流程中使用数据库事务，先锁座再落单，确保并发场景下一致性；取消或退款操作会联动恢复座位状态。",
"5.5 本章小结",
"本章从前端页面、后端接口与运维脚本三方面给出了系统实现过程，验证了设计方案可落地。",

"第6章 系统测试与结果分析",
"6.1 测试环境",
"测试环境：Windows 10、微信开发者工具、Node.js、MySQL。后端地址为 http://127.0.0.1:3000。",
"6.2 功能测试",
"对登录、电影检索、影院场次查询、选座下单、订单流转、评论收藏、管理端操作进行测试，结果均符合预期。",
"6.3 异常与边界测试",
"覆盖未登录访问、参数缺失、重复锁座、重复支付、非法订单删除、网络抖动等场景。系统能够给出明确错误提示并保持数据一致。",
"6.4 性能与稳定性分析",
"在课程项目规模下，系统接口响应与页面交互流畅；通过批量导入脚本可构建千级电影数据；通过事务与状态机保证业务稳定性。",
"6.5 测试结论",
"系统满足毕业设计场景下的主要功能需求，具备较好的可维护性和可扩展性。",

"结论与展望",
"本文完成了基于微信小程序的在线电影购票系统设计与实现，形成了从数据导入、影片展示、场次选座到订单流转的完整业务闭环。系统在并发一致性控制、模块化设计与运维能力方面取得了较好效果。未来可进一步接入真实微信支付、消息通知、推荐算法、缓存与分布式架构优化，提升系统的工程化与商业化能力。",

"参考文献",
"[1] 微信开放文档. 微信小程序开发文档.",
"[2] Express.js Official Documentation.",
"[3] MySQL 8.0 Reference Manual.",
"[4] The Movie Database (TMDB) API Documentation.",
"[5] Fielding R T. Architectural Styles and the Design of Network-based Software Architectures.",
"[6] Martin Fowler. Patterns of Enterprise Application Architecture.",
"[7] IEEE Software Engineering Body of Knowledge.",

"致谢",
"感谢指导教师在课题方向、系统设计和论文写作中的悉心指导；感谢同学在系统测试和联调过程中的帮助；感谢家人对学习与毕业设计工作的支持。"
]

for i, line in enumerate(content):
    pos = start_idx + i
    if pos < len(doc.paragraphs):
        doc.paragraphs[pos].text = line
    else:
        doc.add_paragraph(line)

# 清空后续残留旧内容
end_used = start_idx + len(content)
for pos in range(end_used, len(doc.paragraphs)):
    doc.paragraphs[pos].text = ""

doc.save(out)
print(out)
print('written_lines=', len(content))
