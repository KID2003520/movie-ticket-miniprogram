from docx import Document

src = r"C:\Users\zhu20\Desktop\毕业论文\毕设论文参考模板2026.docx"
out = r"C:\Users\zhu20\Desktop\毕业论文\毕设论文参考模板2026_电影购票系统_完成版.docx"

doc = Document(src)

# 1) 封面与题目相关（按模板现有位置覆盖）
cover_updates = {
    12: "题目：基于微信小程序的在线电影购票系统设计与实现",
    29: "毕业设计题目：基于微信小程序的在线电影购票平台",
    33: "本课题围绕“基于微信小程序的在线电影购票系统”开展设计与实现，构建一个覆盖用户端与管理端的完整票务业务系统。系统实现电影浏览、影院场次查询、在线选座、订单创建/支付/取消/退款、收藏评论、管理端电影管理与用户管理等功能，并支持基于 TMDB 的真实电影数据导入。",
    67: "关键词：在线电影购票；微信小程序；Node.js；Express；MySQL；选座系统",
    72: "Key words: online movie ticketing; WeChat Mini Program; Node.js; Express; MySQL; seat selection"
}

for idx, txt in cover_updates.items():
    if 0 <= idx < len(doc.paragraphs):
        doc.paragraphs[idx].text = txt

# 2) 摘要/英文摘要覆盖
if len(doc.paragraphs) > 66:
    doc.paragraphs[65].text = "摘  要"
    doc.paragraphs[66].text = (
        "随着移动互联网与在线支付的普及，用户对电影票务服务的即时性与便捷性要求不断提高。"
        "本文以“在线电影购票微信小程序”为研究对象，完成了系统需求分析、总体架构设计、数据库设计、功能实现与测试验证。"
        "系统前端采用微信小程序原生框架，后端采用 Node.js + Express，数据层采用 MySQL，形成前后端分离的实现方案。"
        "在功能层面，系统实现了电影列表与详情、影院检索、场次查询、可视化选座、订单流转、收藏评论、管理后台等核心业务。"
        "在关键机制上，系统通过事务化锁座与状态流转控制，保障并发购票场景下的数据一致性。"
        "同时，系统支持 TMDB 批量导入，能够快速构建真实电影数据集，提升系统演示与应用价值。"
        "测试结果表明，该系统具备较好的可用性、稳定性和可扩展性，能够满足毕业设计场景下的业务需求。"
    )
    doc.paragraphs[67].text = "关键词：在线电影购票；微信小程序；Node.js；Express；MySQL；选座系统"

if len(doc.paragraphs) > 72:
    doc.paragraphs[69].text = "ABSTRACT"
    doc.paragraphs[70].text = (
        "With the rapid growth of mobile Internet services, users expect more efficient and convenient movie ticket booking experiences. "
        "This thesis designs and implements an online movie ticketing system based on WeChat Mini Program. "
        "The frontend is built with native mini-program technologies, while the backend adopts Node.js and Express with MySQL as the persistent storage. "
        "The system supports movie browsing, cinema and schedule query, visual seat selection, order creation/payment/cancellation/refund, collection and comment, and admin management. "
        "A transaction-based seat-locking strategy is used to ensure data consistency under concurrent purchase requests. "
        "In addition, TMDB batch import is integrated to provide real movie datasets for deployment and demonstration. "
        "Experimental results show that the system is stable, practical, and extensible for graduation project scenarios."
    )
    doc.paragraphs[71].text = "Key words: online movie ticketing; WeChat Mini Program; Node.js; Express; MySQL; seat selection"

# 3) 从“目录后正文”开始整体重写（保留模板前置页）
start_idx = 74
content = [
"第1章 绪论",
"1.1 研究背景",
"传统电影购票方式存在排队时间长、场次信息不透明、座位选择受限等问题。随着移动互联网发展，用户更倾向于通过手机完成“查影讯—选影院—选座—支付”全流程操作。微信小程序作为轻应用形态，具有无需安装、触达便捷、生态成熟等优势，适合作为在线电影购票系统的业务载体。",
"1.2 研究意义",
"本课题将电影票务核心流程完整落地到微信小程序中，既有助于验证前后端分离架构在真实业务中的可行性，也为后续接入真实支付、推荐算法和运营系统提供了可扩展基础。",
"1.3 研究内容",
"本文围绕在线电影购票系统完成以下工作：需求分析、系统架构设计、数据库建模、核心业务实现、系统测试与结果分析。",
"第2章 关键技术与开发环境",
"2.1 微信小程序技术",
"系统前端使用 WXML、WXSS 和 JavaScript 构建页面与交互，通过页面生命周期和本地缓存机制实现状态管理，结合微信开放能力完成登录与位置相关功能。",
"2.2 Node.js 与 Express",
"后端采用 Express 构建 REST API，统一处理电影、影院、场次、订单、评论、收藏及管理端请求，具备开发效率高、生态完善、易于扩展的特点。",
"2.3 MySQL 数据库",
"系统使用 MySQL 进行业务数据持久化，围绕 movies、cinemas、schedules、seats、orders、users、collections、movie_comments 等表构建票务数据模型。",
"2.4 TMDB 数据接入",
"系统支持通过 TMDB API 批量导入电影数据，结合代理配置可稳定获取真实电影信息，并通过状态重算脚本更新上映状态。",
"第3章 系统需求分析",
"3.1 用户端功能需求",
"用户端需支持：电影浏览与检索、影院与场次查询、在线选座、订单支付与管理、收藏与评论、个人信息维护。",
"3.2 管理端功能需求",
"管理端需支持：电影状态维护、电影删除、用户信息查看、统计数据查看、场次重建与数据导入管理。",
"3.3 非功能需求",
"系统需满足可用性、稳定性、安全性与可维护性要求；并发购票时要保证座位状态一致，避免出现重复售卖。",
"第4章 系统设计",
"4.1 系统总体架构",
"系统采用“微信小程序前端 + Express 后端 + MySQL 数据库”的三层架构，前后端通过统一 API 协议交互，实现业务解耦。",
"4.2 功能模块设计",
"前端包含：首页、电影页、电影详情页、影院页、影院详情页、选座页、订单页、订单详情页、用户中心与管理页面。后端对应提供电影、影院、场次、订单、评论收藏、管理等服务模块。",
"4.3 数据库设计",
"movies 表存储影片基础信息；cinemas 表存储影院信息；schedules 表存储场次；seats 表维护座位实时状态；orders 表记录订单生命周期；collections 与 movie_comments 实现互动功能。",
"4.4 关键流程设计",
"订单创建时先在事务中锁定座位（available→locked），支付成功后更新为 sold；取消或退款时释放座位。该机制确保并发场景下的数据一致性。",
"第5章 系统实现",
"5.1 前端实现",
"小程序通过统一网络层调用后端接口，页面侧根据业务状态进行渲染与交互控制，实现了从电影检索到下单支付的完整用户路径。",
"5.2 后端实现",
"后端提供 /api/movies、/api/cinemas、/api/schedules、/api/orders 等接口。订单相关接口支持创建、取消、模拟支付、退款与删除操作。",
"5.3 数据导入与运维实现",
"系统内置批量导入脚本，可从 TMDB 追加大量电影数据；通过 repairMoviesAndSchedules 脚本批量重建排期与座位，提高演示与部署效率。",
"5.4 管理模块实现",
"管理员可以在小程序管理页面中执行电影上下线、删除影片、查看用户列表和统计信息，支持项目运维与数据管理。",
"第6章 系统测试",
"6.1 测试环境",
"测试环境包括 Windows、微信开发者工具、Node.js、MySQL。后端默认运行在 3000 端口，前端通过配置项连接后端地址。",
"6.2 功能测试",
"对登录、电影浏览、场次查询、选座购票、订单流转、评论收藏、后台管理等核心流程进行测试，结果均符合预期。",
"6.3 异常测试",
"对未登录访问、参数缺失、重复锁座、重复支付、非法订单删除等场景进行测试，系统均能返回明确错误信息并保持数据一致。",
"6.4 测试结果分析",
"系统在主要业务链路上运行稳定，能够支撑毕业设计所需功能展示，并具备后续工程化扩展潜力。",
"结论",
"本文完成了基于微信小程序的在线电影购票系统设计与实现，形成了包含用户端与管理端的完整业务闭环。系统在真实电影数据导入、选座一致性控制和订单状态管理方面实现了可验证的工程效果。",
"展望",
"后续工作可围绕真实支付接入、消息通知、推荐算法、缓存加速与高并发优化展开，进一步提升系统性能与商业可用性。",
"参考文献",
"[1] 微信开放文档. 微信小程序开发文档.",
"[2] Express 官方文档.",
"[3] MySQL 参考手册.",
"[4] TMDB API Documentation.",
"致谢",
"感谢指导教师在课题研究与论文撰写中的指导，感谢同学在系统测试和联调过程中的支持。"
]

# 覆盖现有段落，超出部分追加，剩余旧段落清空
for i, line in enumerate(content):
    pos = start_idx + i
    if pos < len(doc.paragraphs):
        doc.paragraphs[pos].text = line
    else:
        doc.add_paragraph(line)

end_used = start_idx + len(content)
for pos in range(end_used, len(doc.paragraphs)):
    doc.paragraphs[pos].text = ""

doc.save(out)
print(out)
