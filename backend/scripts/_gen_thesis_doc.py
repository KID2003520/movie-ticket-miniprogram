from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn

out = r"C:\Users\zhu20\Desktop\毕业论文\毕设论文_在线电影购票微信小程序_完成版.docx"
d = Document()
style = d.styles['Normal']
style.font.name = '宋体'
style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
style.font.size = Pt(12)

def h(text, lv=1):
    d.add_heading(text, level=lv)

def p(text):
    d.add_paragraph(text)

h('毕业论文', 0)
p('题目：基于微信小程序的在线电影购票系统设计与实现')
p('说明：本稿根据你的项目代码自动整理，可直接在此基础上按学校格式微调。')

h('摘  要', 1)
p('本文围绕“在线电影购票微信小程序”开展系统设计与实现，目标是构建一个可部署、可演示、可扩展的在线票务系统。系统前端采用微信小程序原生框架，后端采用 Node.js + Express + MySQL，并保留云函数兼容能力。')
p('系统实现了电影浏览、影院检索、场次查询、在线选座、订单创建/取消/支付/退款、收藏评论、管理端电影与用户管理、TMDB 批量导入等核心功能。针对高频购票场景，系统在订单创建中采用座位锁定与事务控制，保障并发情况下的数据一致性，避免重复售座。')
p('测试结果表明，系统在主要业务链路上运行稳定，能满足毕业设计场景下的功能与演示需求。')
p('关键词：微信小程序；在线电影购票；Node.js；Express；MySQL；选座系统')

h('ABSTRACT', 1)
p('This thesis presents the design and implementation of an online movie ticketing system based on WeChat Mini Program. The frontend is implemented with native mini-program technologies, while the backend uses Node.js, Express, and MySQL. The system supports complete workflows including movie browsing, cinema schedules, seat selection, order lifecycle management, user interactions, and admin operations.')
p('A transaction-based seat-locking mechanism is introduced to maintain consistency under concurrent requests. Experimental verification shows that the system is stable and practical for graduation project scenarios.')

h('第1章 绪论', 1)
p('随着移动互联网和在线支付的普及，电影票务业务逐步从线下窗口迁移到移动端。微信小程序具备低门槛、无需安装和生态完善的优势，适合构建轻量化票务系统。本文研究并实现一个面向实际业务流程的在线电影购票小程序。')

h('第2章 关键技术', 1)
p('2.1 微信小程序：负责页面渲染、交互逻辑与本地缓存。')
p('2.2 Express 后端：提供统一 REST API，承载电影、影院、排期、订单等业务。')
p('2.3 MySQL 数据库：存储 movies、cinemas、schedules、seats、orders、users 等核心数据。')
p('2.4 TMDB 数据导入：通过脚本批量导入真实电影，支持状态重算与排期重建。')

h('第3章 需求分析', 1)
p('系统需求分为用户端与管理端。用户端包括登录注册、电影检索、选座购票、订单管理、评论收藏；管理端包括影片状态管理、用户查看、排期重建、统计分析。')

h('第4章 系统设计', 1)
p('系统采用前后端分离架构。前端通过统一 API 层访问后端，后端通过服务层和数据层实现业务逻辑。核心数据实体包括电影、影院、场次、座位、订单、评论、收藏。')
p('在一致性设计上，创建订单时将座位从 available 更新为 locked，支付后更新为 sold，取消或退款时回滚为 available。')

h('第5章 系统实现', 1)
p('前端实现：首页、电影页、影院页、影院详情、选座页、订单页、用户中心及管理页面。')
p('后端实现：健康检查、电影列表与详情、影院与场次查询、订单创建/取消/支付/退款、评论收藏、管理端接口。')
p('运维实现：提供 importAllMoviesTmdb、importDiscoverBulk、repairMoviesAndSchedules 等脚本，用于真实数据导入和排期快速恢复。')

h('第6章 系统测试', 1)
p('对登录、浏览、选座、下单、支付、退款、评论收藏、管理端操作等流程进行测试。结果表明核心功能完整可用，异常分支具备明确错误提示。')

h('结论与展望', 1)
p('本文完成了在线电影购票微信小程序的设计与实现。系统已具备完整业务闭环与较好的可维护性。后续可进一步接入真实支付、消息推送、推荐算法及高并发优化，提升工程化水平。')

h('参考文献（示例）', 1)
p('[1] 微信开放文档. 微信小程序开发指南.')
p('[2] Express 官方文档.')
p('[3] MySQL Reference Manual.')
p('[4] TMDB API Documentation.')

h('致谢', 1)
p('感谢指导教师和同学在需求分析、系统联调与论文写作过程中的支持与帮助。')

d.save(out)
print(out)
